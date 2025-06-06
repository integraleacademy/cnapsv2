
from flask import Flask, render_template, request, redirect, send_file
import sqlite3
import os
from docx import Document
from docx.shared import Inches
from io import BytesIO

app = Flask(__name__)
DB_NAME = "cnaps.db"

def init_db():
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("""CREATE TABLE IF NOT EXISTS dossiers (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            nom TEXT,
            prenom TEXT,
            formation TEXT,
            session TEXT,
            statut TEXT DEFAULT 'INCOMPLET',
            commentaire TEXT
        );""")

init_db()

@app.route("/")
def index():
    with sqlite3.connect(DB_NAME) as conn:
        dossiers = conn.execute("SELECT * FROM dossiers").fetchall()
    return render_template("index.html", dossiers=dossiers)

@app.route("/add", methods=["POST"])
def add():
    nom = request.form["nom"]
    prenom = request.form["prenom"]
    formation = request.form["formation"]
    session = request.form["session"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("INSERT INTO dossiers (nom, prenom, formation, session) VALUES (?, ?, ?, ?)",
                     (nom, prenom, formation, session))
    return redirect("/")

@app.route("/update_status/<int:id>", methods=["POST"])
def update_status(id):
    statut = request.form["statut"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET statut = ? WHERE id = ?", (statut, id))
    return redirect("/")

@app.route("/commentaire/<int:id>", methods=["POST"])
def commentaire(id):
    commentaire = request.form["commentaire"]
    with sqlite3.connect(DB_NAME) as conn:
        conn.execute("UPDATE dossiers SET commentaire = ? WHERE id = ?", (commentaire, id))
    return redirect("/")

@app.route("/attestation/<int:id>")
def attestation(id):
    with sqlite3.connect(DB_NAME) as conn:
        dossier = conn.execute("SELECT * FROM dossiers WHERE id = ?", (id,)).fetchone()
    if not dossier:
        return "Dossier introuvable", 404

    nom, prenom, formation, session = dossier[1], dossier[2], dossier[3], dossier[4]
    template = "fichier_aps.docx" if formation == "APS" else "fichier_a3p.docx"
    doc = Document(template)

    for p in doc.paragraphs:
        p.text = p.text.replace("{{nom}}", nom).replace("{{prenom}}", prenom).replace("{{session}}", session)

    doc.add_picture("static/tampon_signature.png", width=Inches(4))
    byte_io = BytesIO()
    doc.save(byte_io)
    byte_io.seek(0)

    return send_file(byte_io, as_attachment=True, download_name=f"attestation_{nom}_{prenom}.docx")

if __name__ == "__main__":
    app.run(debug=True)
